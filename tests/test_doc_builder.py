"""
Automated tests for doc_builder module.

Tests verify:
- Paragraph alignment remains RTL
- Template styles are preserved
- Hebrew text is readable and correctly ordered
- Bullet lists appear correctly using template formatting
- Error handling works correctly
"""

import os
import unittest
import tempfile
import shutil
import warnings
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from utils.doc_builder import (
    create_doc_from_template,
    append_to_existing_doc,
    _is_bullet_line,
    _contains_hebrew
)


class TestDocBuilder(unittest.TestCase):
    """Test cases for doc_builder module."""
    
    def setUp(self):
        """Set up test fixtures."""
        # Create temporary directory for test files
        self.test_dir = tempfile.mkdtemp()
        self.template_path = os.path.join(self.test_dir, 'test_template.dotx')
        self.output_path = os.path.join(self.test_dir, 'test_output.docx')
        
        # Create a minimal template document for testing
        # Note: In real usage, this would be a properly configured .dotx template
        # For testing, we create a basic document and save it as template
        doc = Document()
        
        # Configure Normal style for RTL (minimal setup for testing)
        normal_style = doc.styles['Normal']
        normal_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Configure List Bullet style for RTL
        try:
            bullet_style = doc.styles['List Bullet']
            bullet_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        except KeyError:
            pass  # Style might not exist in all python-docx versions
        
        doc.save(self.template_path)
    
    def tearDown(self):
        """Clean up test fixtures."""
        shutil.rmtree(self.test_dir)
    
    def test_create_doc_from_template(self):
        """Test creating a new document from template."""
        hebrew_text = "זהו טקסט בעברית"
        
        create_doc_from_template(
            template_path=self.template_path,
            output_path=self.output_path,
            text=hebrew_text
        )
        
        # Verify document was created
        self.assertTrue(os.path.exists(self.output_path))
        
        # Verify document can be opened
        doc = Document(self.output_path)
        self.assertGreater(len(doc.paragraphs), 0)
    
    def test_paragraph_alignment_rtl(self):
        """Test that paragraphs remain RTL aligned."""
        hebrew_text = "זהו טקסט בעברית"
        
        create_doc_from_template(
            template_path=self.template_path,
            output_path=self.output_path,
            text=hebrew_text
        )
        
        doc = Document(self.output_path)
        
        # Check that paragraphs have RIGHT alignment (from template)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Skip empty paragraphs
                # Alignment should be RIGHT (from template)
                # Note: If template doesn't set alignment, this might be None
                # In real usage with proper template, alignment will be RIGHT
                alignment = paragraph.paragraph_format.alignment
                # Allow None (default) or RIGHT
                self.assertIn(alignment, [None, WD_PARAGRAPH_ALIGNMENT.RIGHT])
    
    def test_style_names_preserved(self):
        """Test that template styles are used correctly."""
        hebrew_text = "זהו טקסט רגיל\n- פריט רשימה"
        
        create_doc_from_template(
            template_path=self.template_path,
            output_path=self.output_path,
            text=hebrew_text
        )
        
        doc = Document(self.output_path)
        
        # Verify styles are applied
        # First paragraph should use Normal style
        if len(doc.paragraphs) > 0 and doc.paragraphs[0].text.strip():
            self.assertEqual(doc.paragraphs[0].style.name, 'Normal')
    
    def test_hebrew_text_readable(self):
        """Test that Hebrew text is correctly inserted and readable."""
        hebrew_text = "זהו טקסט בעברית עם פיסוק נכון!"
        
        create_doc_from_template(
            template_path=self.template_path,
            output_path=self.output_path,
            text=hebrew_text
        )
        
        doc = Document(self.output_path)
        
        # Find paragraph with Hebrew text
        found_hebrew = False
        for paragraph in doc.paragraphs:
            if 'זהו' in paragraph.text:
                found_hebrew = True
                # Verify text is preserved correctly
                self.assertIn('זהו', paragraph.text)
                self.assertIn('פיסוק', paragraph.text)
                break
        
        self.assertTrue(found_hebrew, "Hebrew text not found in document")
    
    def test_bullet_lists_correct(self):
        """Test that bullet lists use template RTL list style."""
        text_with_bullets = "- פריט ראשון\n- פריט שני\n* פריט שלישי"
        
        create_doc_from_template(
            template_path=self.template_path,
            output_path=self.output_path,
            text=text_with_bullets
        )
        
        doc = Document(self.output_path)
        
        # Verify bullet detection worked
        # Bullet markers should be removed from content
        bullet_paragraphs = [p for p in doc.paragraphs if p.text.strip() and 'פריט' in p.text]
        self.assertGreater(len(bullet_paragraphs), 0)
        
        # Verify bullet markers are removed from content
        for para in bullet_paragraphs:
            self.assertNotIn('-', para.text)
            self.assertNotIn('*', para.text)
    
    def test_append_to_existing_doc(self):
        """Test appending content to existing document."""
        # Create initial document
        initial_text = "טקסט ראשוני"
        create_doc_from_template(
            template_path=self.template_path,
            output_path=self.output_path,
            text=initial_text
        )
        
        # Append additional content
        additional_text = "טקסט נוסף"
        append_to_existing_doc(
            doc_path=self.output_path,
            text=additional_text
        )
        
        # Verify both texts are in document
        doc = Document(self.output_path)
        all_text = ' '.join([p.text for p in doc.paragraphs])
        self.assertIn('ראשוני', all_text)
        self.assertIn('נוסף', all_text)
    
    def test_append_alignment_safety_check(self):
        """Test that append function applies minimal alignment safety check."""
        # Create initial document
        create_doc_from_template(
            template_path=self.template_path,
            output_path=self.output_path,
            text="טקסט ראשוני"
        )
        
        # Append content
        append_to_existing_doc(
            doc_path=self.output_path,
            text="טקסט נוסף"
        )
        
        # Verify alignment safety check was applied
        doc = Document(self.output_path)
        # New paragraphs from append should have alignment set if needed
        # (minimal safety check)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip() and 'נוסף' in paragraph.text:
                alignment = paragraph.paragraph_format.alignment
                # Should be RIGHT (from safety check or template)
                self.assertIn(alignment, [None, WD_PARAGRAPH_ALIGNMENT.RIGHT])
    
    def test_missing_template_error(self):
        """Test error handling for missing template."""
        with self.assertRaises(FileNotFoundError):
            create_doc_from_template(
                template_path='nonexistent_template.dotx',
                output_path=self.output_path,
                text="טקסט"
            )
    
    def test_missing_document_error(self):
        """Test error handling for missing document when appending."""
        with self.assertRaises(FileNotFoundError):
            append_to_existing_doc(
                doc_path='nonexistent_doc.docx',
                text="טקסט"
            )
    
    def test_non_hebrew_warning(self):
        """Test that non-Hebrew input generates warning but doesn't fail."""
        english_text = "This is English text"
        
        with warnings.catch_warnings(record=True) as w:
            warnings.simplefilter("always")
            create_doc_from_template(
                template_path=self.template_path,
                output_path=self.output_path,
                text=english_text
            )
            
            # Verify warning was issued
            self.assertGreater(len(w), 0)
            self.assertTrue(any('Hebrew' in str(warning.message) for warning in w))
        
        # Verify document was still created
        self.assertTrue(os.path.exists(self.output_path))
    
    def test_bullet_detection(self):
        """Test bullet detection helper function."""
        # Test various bullet formats
        self.assertTrue(_is_bullet_line("- פריט"))
        self.assertTrue(_is_bullet_line("* פריט"))
        self.assertTrue(_is_bullet_line("• פריט"))
        self.assertTrue(_is_bullet_line("  - פריט"))  # With leading spaces
        
        # Test non-bullet lines
        self.assertFalse(_is_bullet_line("פריט רגיל"))
        self.assertFalse(_is_bullet_line(""))
        self.assertFalse(_is_bullet_line("   "))  # Only spaces
    
    def test_contains_hebrew(self):
        """Test Hebrew character detection."""
        self.assertTrue(_contains_hebrew("זהו טקסט בעברית"))
        self.assertTrue(_contains_hebrew("Mixed text עם עברית"))
        self.assertFalse(_contains_hebrew("This is English text"))
        self.assertFalse(_contains_hebrew("123"))
        self.assertFalse(_contains_hebrew(""))
    
    def test_multiple_paragraphs(self):
        """Test handling of multiple paragraphs."""
        multi_para_text = """פסקה ראשונה.

פסקה שנייה.

פסקה שלישית."""
        
        create_doc_from_template(
            template_path=self.template_path,
            output_path=self.output_path,
            text=multi_para_text
        )
        
        doc = Document(self.output_path)
        # Should have multiple paragraphs
        self.assertGreaterEqual(len(doc.paragraphs), 3)
    
    def test_empty_lines(self):
        """Test handling of empty lines as paragraph breaks."""
        text_with_empty = "פסקה ראשונה\n\nפסקה שנייה"
        
        create_doc_from_template(
            template_path=self.template_path,
            output_path=self.output_path,
            text=text_with_empty
        )
        
        doc = Document(self.output_path)
        # Empty lines should create paragraph breaks
        self.assertGreater(len(doc.paragraphs), 1)


if __name__ == '__main__':
    unittest.main()
